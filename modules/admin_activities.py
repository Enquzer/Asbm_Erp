from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file
from flask_login import login_required
from datetime import datetime
from modules.models import db, Bill, FoodFuelRecord, SecurityIncident, PettyCash, ProjectFunding, PropertyItem, AdminLetter, DutyStation, Employee, Project, User
import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from openpyxl import Workbook

admin_activities_bp = Blueprint('admin_activities', __name__)

@admin_activities_bp.route('/admin_activities', methods=['GET', 'POST'])
@login_required
def admin_activities():
    duty_stations = DutyStation.query.all()
    employees = Employee.query.all()
    projects = Project.query.all()
    bills = Bill.query.all()
    food_fuel_records = FoodFuelRecord.query.all()
    security_incidents = SecurityIncident.query.all()
    petty_cash_requests = PettyCash.query.all()
    project_funding = ProjectFunding.query.all()
    property_items = PropertyItem.query.all()
    admin_letters = AdminLetter.query.all()

    if request.method == 'POST':
        section = request.form.get('section')
        
        # Bills
        if section == 'bills':
            bill = Bill(
                bill_number=request.form['bill_number'],
                receipt_number=request.form['receipt_number'],
                bill_type=request.form['bill_type'],
                description=request.form['description'] if request.form['bill_type'] == 'Other' else request.form['bill_type'],
                amount=float(request.form['amount']),
                due_date=datetime.strptime(request.form['due_date'], '%Y-%m-%d').date(),
                duty_station_id=int(request.form['duty_station_id']),
                status='Pending'
            )
            db.session.add(bill)
            db.session.commit()
            flash('Bill added successfully!', 'success')

        # Food & Fuel
        elif section == 'food_fuel':
            payee_id = request.form.get('payee_id')
            payee_name = request.form.get('payee_name')
            record = FoodFuelRecord(
                type=request.form['type'],
                description=request.form['description'],
                quantity=float(request.form['quantity']),
                cost=float(request.form['cost']),
                date=datetime.strptime(request.form['date'], '%Y-%m-%d').date(),
                duty_station_id=int(request.form['duty_station_id']),
                payee_id=int(payee_id) if payee_id else None,
                payee_name=payee_name if not payee_id else None
            )
            db.session.add(record)
            db.session.commit()
            flash('Food/Fuel record added successfully!', 'success')

        # Security Incidents
        elif section == 'security':
            incident = SecurityIncident(
                incident_type=request.form['incident_type'],
                description=request.form['description'],
                location=request.form['location'],
                reported_date=datetime.strptime(request.form['reported_date'], '%Y-%m-%d').date(),
                duty_station_id=int(request.form['duty_station_id']),
                status='Open'
            )
            db.session.add(incident)
            db.session.commit()
            flash('Security incident added successfully!', 'success')

        # Petty Cash
        elif section == 'petty_cash':
            employee_id = int(request.form['employee_id'])
            employee = Employee.query.get(employee_id)
            petty_cash = PettyCash(
                description=request.form['description'],
                amount=float(request.form['amount']),
                request_date=datetime.strptime(request.form['request_date'], '%Y-%m-%d').date(),
                employee_id=employee_id,
                employee_title=employee.title,
                reason=request.form['reason'],
                duty_station_id=int(request.form['duty_station_id']),
                status='Pending'
            )
            db.session.add(petty_cash)
            db.session.commit()
            flash('Petty cash request added successfully!', 'success')

        # Project Funding
        elif section == 'project_funding':
            project_id = request.form.get('project_id')
            new_project_name = request.form.get('new_project_name')
            if new_project_name:
                new_project = Project(
                    name=new_project_name,
                    start_date=datetime.strptime(request.form['start_date'], '%Y-%m-%d').date() if request.form['start_date'] else datetime.now().date(),
                    end_date=datetime.strptime(request.form['end_date'], '%Y-%m-%d').date() if request.form['end_date'] else None,
                    status='Active'
                )
                db.session.add(new_project)
                db.session.commit()
                project_id = new_project.id
            funding = ProjectFunding(
                project_id=int(project_id),
                amount=float(request.form['amount']),
                funding_date=datetime.strptime(request.form['funding_date'], '%Y-%m-%d').date(),
                start_date=datetime.strptime(request.form['start_date'], '%Y-%m-%d').date() if request.form['start_date'] else None,
                end_date=datetime.strptime(request.form['end_date'], '%Y-%m-%d').date() if request.form['end_date'] else None,
                source=request.form['source'],
                duty_station_id=int(request.form['duty_station_id'])
            )
            db.session.add(funding)
            db.session.commit()
            flash('Funding added successfully!', 'success')

        # Property Items
        elif section == 'property_items':
            employee_id = request.form.get('employee_id')
            item = PropertyItem(
                item_code=request.form['item_code'],
                item_type=request.form['item_type'],
                description=request.form['description'],
                assigned_date=datetime.strptime(request.form['assigned_date'], '%Y-%m-%d').date() if request.form['assigned_date'] else None,
                employee_id=int(employee_id) if employee_id else None,
                duty_station_id=int(request.form['duty_station_id']),
                status='In Use'
            )
            db.session.add(item)
            db.session.commit()
            flash('Property item added successfully!', 'success')

        # Admin Letters
        elif section == 'admin_letters':
            letter = AdminLetter(
                letter_type=request.form['letter_type'],
                recipient=request.form['recipient'],
                subject=request.form['subject'],
                content=request.form['content'],
                duty_station_id=int(request.form['duty_station_id'])
            )
            db.session.add(letter)
            db.session.commit()
            flash('Letter added successfully!', 'success')

        return redirect(url_for('admin_activities.admin_activities'))

    return render_template('admin_activities.html', bills=bills, food_fuel_records=food_fuel_records,
                           security_incidents=security_incidents, petty_cash_requests=petty_cash_requests,
                           project_funding=project_funding, property_items=property_items, admin_letters=admin_letters,
                           duty_stations=duty_stations, employees=employees, projects=projects)

# Edit Routes
@admin_activities_bp.route('/edit_bill/<int:id>', methods=['POST'])
@login_required
def edit_bill(id):
    bill = Bill.query.get_or_404(id)
    bill.bill_number = request.form['bill_number']
    bill.receipt_number = request.form['receipt_number']
    bill.bill_type = request.form['bill_type']
    bill.description = request.form['description'] if request.form['bill_type'] == 'Other' else request.form['bill_type']
    bill.amount = float(request.form['amount'])
    bill.due_date = datetime.strptime(request.form['due_date'], '%Y-%m-%d').date()
    bill.duty_station_id = int(request.form['duty_station_id'])
    db.session.commit()
    flash('Bill updated successfully!', 'success')
    return redirect(url_for('admin_activities.admin_activities'))

@admin_activities_bp.route('/edit_food_fuel/<int:id>', methods=['POST'])
@login_required
def edit_food_fuel(id):
    record = FoodFuelRecord.query.get_or_404(id)
    record.type = request.form['type']
    record.description = request.form['description']
    record.quantity = float(request.form['quantity'])
    record.cost = float(request.form['cost'])
    record.date = datetime.strptime(request.form['date'], '%Y-%m-%d').date()
    record.duty_station_id = int(request.form['duty_station_id'])
    record.payee_id = int(request.form['payee_id']) if request.form['payee_id'] else None
    record.payee_name = request.form['payee_name'] if not request.form['payee_id'] else None
    db.session.commit()
    flash('Food/Fuel record updated successfully!', 'success')
    return redirect(url_for('admin_activities.admin_activities'))

# Delete Routes
@admin_activities_bp.route('/delete_bill/<int:id>')
@login_required
def delete_bill(id):
    bill = Bill.query.get_or_404(id)
    db.session.delete(bill)
    db.session.commit()
    flash('Bill deleted successfully!', 'success')
    return redirect(url_for('admin_activities.admin_activities'))

@admin_activities_bp.route('/delete_food_fuel/<int:id>')
@login_required
def delete_food_fuel(id):
    record = FoodFuelRecord.query.get_or_404(id)
    db.session.delete(record)
    db.session.commit()
    flash('Food/Fuel record deleted successfully!', 'success')
    return redirect(url_for('admin_activities.admin_activities'))

# Overview Routes
@admin_activities_bp.route('/bills_overview', methods=['POST'])
@login_required
def bills_overview():
    duty_stations = DutyStation.query.all()
    start_date = request.form.get('start_date')
    end_date = request.form.get('end_date')
    overview = {}
    for ds in duty_stations:
        query = Bill.query.filter_by(duty_station_id=ds.id)
        if start_date and end_date:
            query = query.filter(Bill.due_date.between(start_date, end_date))
        total = sum(bill.amount for bill in query.all())
        overview[ds.name] = total
    return render_template('bills_overview.html', overview=overview)

@admin_activities_bp.route('/petty_cash_overview', methods=['POST'])
@login_required
def petty_cash_overview():
    duty_stations = DutyStation.query.all()
    start_date = request.form.get('start_date')
    end_date = request.form.get('end_date')
    overview = {}
    for ds in duty_stations:
        query = PettyCash.query.filter_by(duty_station_id=ds.id, status='Approved')
        if start_date and end_date:
            query = query.filter(PettyCash.request_date.between(start_date, end_date))
        total = sum(req.amount for req in query.all())
        overview[ds.name] = total
    return render_template('petty_cash_overview.html', overview=overview)

@admin_activities_bp.route('/project_funding_overview', methods=['POST'])
@login_required
def project_funding_overview():
    duty_stations = DutyStation.query.all()
    start_date = request.form.get('start_date')
    end_date = request.form.get('end_date')
    overview = {}
    for ds in duty_stations:
        query = ProjectFunding.query.filter_by(duty_station_id=ds.id)
        if start_date and end_date:
            query = query.filter(ProjectFunding.funding_date.between(start_date, end_date))
        projects_funds = {}
        for funding in query.all():
            project_name = funding.project.name if funding.project else 'N/A'
            projects_funds[project_name] = projects_funds.get(project_name, 0) + funding.amount
        overview[ds.name] = projects_funds
    return render_template('project_funding_overview.html', overview=overview)

@admin_activities_bp.route('/property_items_overview', methods=['POST'])
@login_required
def property_items_overview():
    duty_stations = DutyStation.query.all()
    overview = {}
    for ds in duty_stations:
        items = PropertyItem.query.filter_by(duty_station_id=ds.id).all()
        type_count = {}
        for item in items:
            type_count[item.item_type] = type_count.get(item.item_type, 0) + 1
        overview[ds.name] = type_count
    return render_template('property_items_overview.html', overview=overview)

@admin_activities_bp.route('/admin_letters_overview', methods=['POST'])
@login_required
def admin_letters_overview():
    duty_stations = DutyStation.query.all()
    start_date = request.form.get('start_date')
    end_date = request.form.get('end_date')
    overview = {}
    for ds in duty_stations:
        query = AdminLetter.query.filter_by(duty_station_id=ds.id)
        if start_date and end_date:
            query = query.filter(AdminLetter.created_at.between(start_date, end_date))
        letters = query.all()
        internal = sum(1 for l in letters if l.letter_type == 'Internal')
        external = sum(1 for l in letters if l.letter_type == 'External')
        overview[ds.name] = {'Internal': internal, 'External': external}
    return render_template('admin_letters_overview.html', overview=overview)

# Export Routes
@admin_activities_bp.route('/export_pdf/<section>')
@login_required
def export_pdf(section):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    data = []

    if section == 'bills':
        data.append(['Bill #', 'Receipt #', 'Type', 'Amount', 'Due Date', 'Duty Station', 'Status'])
        for bill in Bill.query.all():
            data.append([bill.bill_number, bill.receipt_number, bill.bill_type, bill.amount, bill.due_date, bill.duty_station.name, bill.status])
    # Add similar logic for other sections

    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"{section}_report.pdf", mimetype='application/pdf')

@admin_activities_bp.route('/export_word/<section>')
@login_required
def export_word(section):
    doc = Document()
    doc.add_heading(f'{section.capitalize()} Report', 0)
    table = doc.add_table(rows=1, cols=7 if section == 'bills' else 6)
    hdr_cells = table.rows[0].cells

    if section == 'bills':
        hdr_cells[0].text = 'Bill #'
        hdr_cells[1].text = 'Receipt #'
        hdr_cells[2].text = 'Type'
        hdr_cells[3].text = 'Amount'
        hdr_cells[4].text = 'Due Date'
        hdr_cells[5].text = 'Duty Station'
        hdr_cells[6].text = 'Status'
        for bill in Bill.query.all():
            row_cells = table.add_row().cells
            row_cells[0].text = bill.bill_number
            row_cells[1].text = bill.receipt_number or ''
            row_cells[2].text = bill.bill_type
            row_cells[3].text = str(bill.amount)
            row_cells[4].text = bill.due_date.strftime('%Y-%m-%d')
            row_cells[5].text = bill.duty_station.name
            row_cells[6].text = bill.status
    # Add similar logic for other sections

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"{section}_report.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@admin_activities_bp.route('/export_excel/<section>')
@login_required
def export_excel(section):
    wb = Workbook()
    ws = wb.active
    if section == 'project_funding':
        ws.title = "Project Funding Report"
        ws.append(["Project", "Amount", "Funding Date", "Source", "Duty Station"])
        for funding in ProjectFunding.query.all():
            ws.append([
                funding.project.name if funding.project else 'N/A',
                funding.amount,
                funding.funding_date.strftime('%Y-%m-%d'),
                funding.source,
                funding.duty_station.name
            ])
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"{section}_report.xlsx", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')